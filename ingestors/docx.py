#!/usr/bin/env python3
"""
ingestor: docx — extrai o TEXTO de um .docx recebido no stdin e devolve texto puro no stdout.

Pinagem (contrato do ingestor): stdin = bytes do arquivo · stdout = texto extraído ·
exit 0 = ok, !=0 + motivo no stderr = rejeitado. python3 RAW, sem venv.
Dependência (não-stdlib): python-docx — no Debian/Ubuntu/Mint: `apt install python3-docx`.

Saída: parágrafos do corpo, um por linha; tabelas viram linhas com células separadas por
TAB (o docx é família narrativa/documento — não força CSV como o xlsx). Só .docx (OOXML);
o .doc binário antigo não é suportado por esta engine.
"""
import sys
import os
# uniforme com os demais drivers: remove o próprio diretório do sys.path (anti-sombreamento
# da stdlib — ver README).
_here = os.path.dirname(os.path.abspath(__file__))
sys.path[:] = [p for p in sys.path if p and os.path.abspath(p) != _here]
import io


def main():
    if sys.stdin.isatty():
        sys.stderr.write("uso: python3 docx.py < arquivo.docx   (a fonte vem pelo stdin)\n")
        return 2
    try:
        import docx   # python-docx
    except ImportError:
        sys.stderr.write("docx: dependência ausente: python-docx (apt install python3-docx)\n")
        return 3
    data = sys.stdin.buffer.read()
    if not data:
        sys.stderr.write("docx: entrada vazia\n")
        return 1
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        sys.stderr.write(f"docx: arquivo inválido (.doc antigo não é suportado; use .docx): {e}\n")
        return 1
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    n_paras = len(lines)
    n_tab = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
                n_tab += 1
    if not lines:
        sys.stderr.write("docx: nenhum texto no corpo\n")
        return 1
    sys.stdout.write("\n".join(lines))
    sys.stdout.write("\n")
    sys.stderr.write(f"docx: {n_paras} parágrafo(s) + {n_tab} linha(s) de tabela\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
